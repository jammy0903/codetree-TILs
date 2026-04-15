from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 페이지 여백 ───────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ─── 색상 팔레트 ───────────────────────────────────────────
C_GREEN_BG  = RGBColor(0xC6, 0xEF, 0xCE)   # 출제됨 배경
C_GREEN_TXT = RGBColor(0x27, 0x6F, 0x38)   # 출제됨 글자
C_YELLOW_BG = RGBColor(0xFF, 0xEB, 0x9C)   # 약출제 배경
C_YELLOW_TXT= RGBColor(0x9C, 0x65, 0x00)   # 약출제 글자
C_RED_BG    = RGBColor(0xFF, 0xC7, 0xCE)   # 미출제 배경
C_RED_TXT   = RGBColor(0x9C, 0x00, 0x06)   # 미출제 글자
C_HEADER_BG = RGBColor(0x26, 0x3D, 0x5E)   # 표 헤더 배경
C_HEADER_TXT= RGBColor(0xFF, 0xFF, 0xFF)   # 표 헤더 글자
C_ROW_ALT   = RGBColor(0xF2, 0xF2, 0xF2)   # 짝수행 배경
C_TITLE_TXT = RGBColor(0x1F, 0x39, 0x64)   # 제목 글자


# ─── 헬퍼: 셀 배경색 ──────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


# ─── 헬퍼: 셀 테두리 ──────────────────────────────────────
def set_cell_border(cell, color='BBBBBB', sz=4):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


# ─── 헬퍼: 셀 안쪽 여백 ───────────────────────────────────
def set_cell_margin(cell, top=60, bottom=60, left=100, right=100):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'),    str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)


# ─── 헬퍼: 상태 → 색상 ────────────────────────────────────
def status_colors(status: str):
    if '출제됨' in status and '약' not in status and '미' not in status:
        return C_GREEN_BG, C_GREEN_TXT
    elif '약출제' in status:
        return C_YELLOW_BG, C_YELLOW_TXT
    elif '미출제' in status:
        return C_RED_BG, C_RED_TXT
    return None, None


# ─── 헬퍼: 단락 스타일 ────────────────────────────────────
def styled_para(text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


# ─── 헬퍼: 표 만들기 ──────────────────────────────────────
def make_table(headers, rows):
    col_n = len(headers)
    tbl   = doc.add_table(rows=1 + len(rows), cols=col_n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 열 너비 (항목 / 상태 / 비고)
    widths = [Cm(7.0), Cm(2.5), Cm(6.5)]
    if col_n == 2:
        widths = [Cm(7.0), Cm(9.0)]

    # 헤더 행
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.width = widths[i] if i < len(widths) else Cm(5)
        set_cell_bg(cell, C_HEADER_BG)
        set_cell_border(cell, 'FFFFFF', 6)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold       = True
        run.font.size       = Pt(9)
        run.font.color.rgb  = C_HEADER_TXT

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        tr   = tbl.rows[r_idx + 1]
        is_even = (r_idx % 2 == 1)

        for c_idx, cell_text in enumerate(row_data):
            cell = tr.cells[c_idx]
            cell.width = widths[c_idx] if c_idx < len(widths) else Cm(5)
            set_cell_border(cell, 'DDDDDD', 4)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p   = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)

            # 상태 열(index 1)에 색상 적용
            if c_idx == 1 and col_n == 3:
                bg, fg = status_colors(cell_text)
                if bg:
                    set_cell_bg(cell, bg)
                    run.font.color.rgb = fg
                    run.font.bold      = True
                elif is_even:
                    set_cell_bg(cell, C_ROW_ALT)
            elif col_n == 2 and c_idx == 1:
                # 2열 표 (섹션 8)
                bg, fg = status_colors(cell_text)
                if bg:
                    set_cell_bg(cell, bg)
                    run.font.color.rgb = fg
                    run.font.bold = True
                elif is_even:
                    set_cell_bg(cell, C_ROW_ALT)
            else:
                if is_even:
                    set_cell_bg(cell, C_ROW_ALT)

    return tbl


# ─── 제목 ──────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('정보처리기사 실기 출제현황 최종판')
run.font.size  = Pt(18)
run.font.bold  = True
run.font.color.rgb = C_TITLE_TXT
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(6)

# ─── 판정 기준 박스 ────────────────────────────────────────
legend_items = [
    ('출제됨', C_GREEN_BG, C_GREEN_TXT,   '정답 칸에 직접 등장'),
    ('약출제', C_YELLOW_BG, C_YELLOW_TXT, '문제/보기/요약/해설에서만 확인'),
    ('미출제', C_RED_BG,    C_RED_TXT,    '보유 기출 파일 기준 직접 확인되지 않음'),
]
leg_tbl = doc.add_table(rows=1, cols=3)
leg_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, bg, fg, desc) in enumerate(legend_items):
    cell = leg_tbl.rows[0].cells[i]
    cell.width = Cm(5.3)
    set_cell_bg(cell, bg)
    set_cell_border(cell, 'AAAAAA', 4)
    set_cell_margin(cell, 80, 80, 120, 120)
    p    = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(label)
    run1.font.bold      = True
    run1.font.size      = Pt(9)
    run1.font.color.rgb = fg
    run2 = p.add_run(f'\n{desc}')
    run2.font.size      = Pt(8)
    run2.font.color.rgb = fg

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════
# 섹션 헬퍼
# ════════════════════════════════════════════════════════════
def add_section(title, level=1):
    p = doc.add_heading(title, level=level)
    run = p.runs[0] if p.runs else p.add_run(title)
    run.font.color.rgb = C_TITLE_TXT
    if level == 1:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)


def add_table_section(sub_title, headers, rows):
    add_section(sub_title, level=2)
    make_table(headers, rows)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════════
# 1. 소프트웨어공학
# ════════════════════════════════════════════════════════════
add_section('1. 소프트웨어공학')

add_table_section('1-1. 디자인 패턴', ['항목', '상태', '비고'], [
    ['Observer',            '출제됨', '직접 답'],
    ['Proxy',               '출제됨', '직접 답'],
    ['Factory Method',      '출제됨', '직접 답'],
    ['Singleton',           '출제됨', '직접 답'],
    ['Visitor',             '출제됨', '직접 답'],
    ['Abstract Factory',    '출제됨', '직접 답'],
    ['Iterator',            '출제됨', '직접 답'],
    ['Adapter',             '출제됨', '직접 답'],
    ['Bridge',              '출제됨', '직접 답'],
    ['행위 패턴',           '출제됨', '직접 답'],
    ['생성 패턴',           '약출제', '분류 문제 문맥'],
    ['구조 패턴',           '약출제', '분류 문제 문맥'],
    ['Command',             '약출제', '보기/분류 문맥'],
    ['Chain of Responsibility','약출제','보기/분류 문맥'],
    ['Strategy',            '미출제', ''],
    ['Decorator',           '미출제', ''],
    ['Composite',           '미출제', ''],
    ['Facade',              '미출제', ''],
    ['Builder',             '미출제', ''],
    ['Prototype',           '미출제', ''],
])

add_table_section('1-2. 결합도 / 응집도', ['항목', '상태', '비고'], [
    ['내용 결합도',   '출제됨', '직접 답'],
    ['스탬프 결합도', '출제됨', '직접 답'],
    ['제어 결합도',   '출제됨', '직접 답'],
    ['공통 결합도',   '출제됨', '직접 답'],
    ['자료 결합도',   '미출제', ''],
    ['외부 결합도',   '미출제', ''],
    ['기능적 응집도', '출제됨', '직접 답'],
    ['순차적 응집도', '출제됨', '직접 답'],
    ['교환적 응집도', '출제됨', '직접 답'],
    ['절차적 응집도', '출제됨', '직접 답'],
    ['시간적 응집도', '출제됨', '순서 답안에 직접 포함'],
    ['논리적 응집도', '미출제', ''],
    ['우연적 응집도', '출제됨', '순서 답안에 직접 포함'],
])

add_table_section('1-3. 테스트 / 커버리지', ['항목', '상태', '비고'], [
    ['블랙박스 테스트',         '출제됨', '직접 답'],
    ['화이트박스 테스트',       '출제됨', '직접 답'],
    ['동등분할',               '출제됨', '직접 답'],
    ['경계값 분석',             '출제됨', '직접 답'],
    ['원인-효과 그래프',        '출제됨', '직접 답'],
    ['샘플링 오라클',           '출제됨', '직접 답'],
    ['문장 커버리지',           '출제됨', '직접 답'],
    ['분기 커버리지',           '출제됨', '직접 답'],
    ['결정 커버리지',           '출제됨', '직접 답'],
    ['조건 커버리지',           '출제됨', '직접 답'],
    ['변경 조건/결정 커버리지', '출제됨', '직접 답'],
    ['단위 테스트',             '출제됨', '직접 답'],
    ['통합 테스트',             '출제됨', '직접 답'],
    ['시스템 테스트',           '출제됨', '직접 답'],
    ['인수 테스트',             '출제됨', '직접 답'],
    ['스텁',                   '출제됨', '직접 답'],
    ['드라이버',               '출제됨', '직접 답'],
    ['테스트 조건',             '출제됨', '직접 답'],
    ['테스트 데이터',           '출제됨', '직접 답'],
    ['예상 결과',               '출제됨', '직접 답'],
    ['기본 경로 커버리지',      '미출제', ''],
    ['다중 조건 커버리지',      '미출제', ''],
    ['상태 전이 테스트',        '미출제', ''],
    ['결정 테이블 테스트',      '미출제', ''],
    ['비교 테스트',             '미출제', ''],
    ['분류 트리 테스트',        '미출제', ''],
    ['페어와이즈 테스트',       '미출제', ''],
    ['유스케이스 테스트',       '미출제', ''],
])

add_table_section('1-4. UML / 모델링', ['항목', '상태', '비고'], [
    ['패키지 다이어그램',       '출제됨', '직접 답'],
    ['클래스 다이어그램',       '출제됨', '직접 답'],
    ['연관 관계',               '출제됨', '직접 답'],
    ['일반화 관계',             '출제됨', '직접 답'],
    ['의존 관계',               '출제됨', '직접 답'],
    ['집약(Aggregation)',       '출제됨', '직접 답'],
    ['Function Modeling',       '출제됨', '직접 답'],
    ['Dynamic Modeling',        '출제됨', '직접 답'],
    ['Object Modeling',         '출제됨', '직접 답'],
    ['합성(Composition)',       '미출제', ''],
    ['객체 다이어그램',         '미출제', ''],
    ['컴포넌트 다이어그램',     '미출제', ''],
    ['배치 다이어그램',         '미출제', ''],
    ['복합체 구조 다이어그램',  '미출제', ''],
    ['유스케이스 다이어그램',   '미출제', ''],
    ['활동 다이어그램',         '미출제', ''],
    ['시퀀스 다이어그램',       '미출제', ''],
    ['상태 다이어그램',         '약출제', '럼바우 문맥'],
])

add_table_section('1-5. 기타 SE', ['항목', '상태', '비고'], [
    ['살충제 패러독스',   '출제됨', '직접 답'],
    ['리팩토링',         '출제됨', '직접 답'],
    ['헤더(Header)',     '출제됨', '직접 답'],
    ['LoC 기법',        '출제됨', '직접 답'],
    ['처리량',           '출제됨', '직접 답'],
    ['응답시간',         '출제됨', '직접 답'],
    ['경과시간',         '출제됨', '직접 답'],
    ['기능적 요구사항',  '출제됨', '직접 답'],
    ['비기능적 요구사항','출제됨', '직접 답'],
    ['애자일 방법론',    '출제됨', '직접 답'],
    ['형상관리',         '출제됨', '직접 답'],
    ['형상 통제',        '출제됨', '직접 답'],
    ['정적 분석 도구',   '출제됨', '직접 답'],
    ['정적 분석',        '출제됨', '직접 답'],
    ['동적 분석',        '출제됨', '직접 답'],
    ['JUnit',           '출제됨', '직접 답'],
    ['GUI',             '출제됨', '직접 답'],
    ['NUI',             '출제됨', '직접 답'],
    ['UX',              '출제됨', '직접 답'],
    ['UI',              '출제됨', '직접 답'],
    ['유효성',           '출제됨', '직접 답'],
    ['직관성',           '출제됨', '직접 답'],
    ['Fan-in',          '출제됨', '직접 답'],
    ['Fan-out',         '출제됨', '직접 답'],
    ['템퍼프루핑',       '출제됨', '직접 답'],
    ['ISP',             '출제됨', '직접 답'],
    ['알파 테스트',      '출제됨', '직접 답'],
    ['베타 테스트',      '출제됨', '직접 답'],
    ['회귀 테스트',      '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 2. SQL
# ════════════════════════════════════════════════════════════
add_section('2. SQL')

add_table_section('2-1. 문장 / 키워드', ['항목', '상태', '비고'], [
    ['SELECT',      '출제됨', '직접 답'],
    ['FROM',        '출제됨', '직접 답'],
    ['WHERE',       '출제됨', '직접 답'],
    ['GROUP BY',    '출제됨', '직접 답'],
    ['HAVING',      '출제됨', '직접 답'],
    ['ORDER BY',    '출제됨', '직접 답'],
    ['DESC',        '출제됨', '직접 답'],
    ['INSERT INTO', '출제됨', '직접 답'],
    ['VALUES',      '출제됨', '직접 답'],
    ['UPDATE',      '출제됨', '직접 답'],
    ['SET',         '출제됨', '직접 답'],
    ['DELETE FROM', '출제됨', '직접 답'],
    ['CREATE INDEX','출제됨', '직접 답'],
    ['ALTER TABLE', '출제됨', '직접 답'],
    ['ADD',         '출제됨', '직접 답'],
    ['DROP VIEW',   '출제됨', '직접 답'],
    ['CASCADE',     '출제됨', '직접 답'],
    ['GRANT',       '출제됨', '직접 답'],
    ['REVOKE',      '미출제', ''],
    ['ASC',         '미출제', ''],
])

add_table_section('2-2. 조건 / 연산 / 함수', ['항목', '상태', '비고'], [
    ['COUNT(*)',        '출제됨', '직접 답'],
    ['COUNT(컬럼)',     '출제됨', '직접 답'],
    ['COUNT(DISTINCT)', '출제됨', '직접 답'],
    ['MIN',             '출제됨', '직접 답'],
    ['MAX',             '출제됨', '직접 답'],
    ['AVG',             '출제됨', '직접 답'],
    ['DISTINCT',        '출제됨', '직접 답'],
    ['IN',              '출제됨', '직접 답'],
    ['LIKE',            '출제됨', '직접 답'],
    ['ALL',             '출제됨', '직접 답'],
    ['AND',             '출제됨', '직접 답'],
    ['OR',              '출제됨', '직접 답'],
    ['ON',              '출제됨', '직접 답'],
    ['AS',              '출제됨', '직접 답'],
])

add_table_section('2-3. 조인 / 집합', ['항목', '상태', '비고'], [
    ['CROSS JOIN',       '출제됨', '직접 답'],
    ['UNION',            '출제됨', '직접 답'],
    ['INNER JOIN',       '미출제', ''],
    ['NATURAL JOIN',     '미출제', ''],
    ['LEFT OUTER JOIN',  '미출제', ''],
    ['RIGHT OUTER JOIN', '미출제', ''],
    ['INTERSECT',        '미출제', ''],
    ['MINUS',            '미출제', ''],
])

add_table_section('2-4. 직접 답안 형태', ['항목', '상태', '비고'], [
    ['SELECT 학번, 이름 FROM 학생 WHERE 학년 IN (3,4);',            '출제됨', '직접 답'],
    ['DELETE FROM 학생 WHERE 이름=\'민수\';',                       '출제됨', '직접 답'],
    ['CREATE INDEX IDX_NAME ON 학생(NAME);',                        '출제됨', '직접 답'],
    ['ALTER TABLE 학생 ADD 주소 VARCHAR(20);',                      '출제됨', '직접 답'],
    ['SELECT 과목이름, MIN(점수)... HAVING AVG(점수)>=90',          '출제됨', '직접 답'],
    ['INSERT INTO 학생(... ) VALUES(... )',                          '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 3. 데이터베이스
# ════════════════════════════════════════════════════════════
add_section('3. 데이터베이스')

add_table_section('3-1. 기본 용어', ['항목', '상태', '비고'], [
    ['튜플',            '출제됨', '직접 답'],
    ['릴레이션 인스턴스','출제됨', '직접 답'],
    ['카디널리티',      '출제됨', '직접 답'],
    ['차수(Degree)',    '출제됨', '직접 답'],
    ['속성(Attribute)', '출제됨', '직접 답'],
    ['도메인',          '출제됨', '직접 답'],
    ['DB 스키마',       '출제됨', '직접 답'],
    ['데이터 마이닝',   '출제됨', '직접 답'],
    ['인덱스(색인)',     '출제됨', '직접 답'],
])

add_table_section('3-2. 키 / 무결성', ['항목', '상태', '비고'], [
    ['후보키',       '출제됨', '직접 답'],
    ['외래키',       '출제됨', '직접 답'],
    ['대체키',       '출제됨', '직접 답'],
    ['슈퍼키',       '출제됨', '직접 답'],
    ['기본키',       '미출제', ''],
    ['개체 무결성',  '출제됨', '직접 답'],
    ['참조 무결성',  '출제됨', '직접 답'],
    ['도메인 무결성','출제됨', '직접 답'],
])

add_table_section('3-3. 정규화 / 이상현상', ['항목', '상태', '비고'], [
    ['2NF',                 '출제됨', '직접 답'],
    ['3NF',                 '출제됨', '직접 답'],
    ['1NF',                 '미출제', ''],
    ['BCNF',                '미출제', ''],
    ['4NF',                 '미출제', ''],
    ['5NF',                 '미출제', ''],
    ['반정규화',             '출제됨', '직접 답'],
    ['삽입 이상',           '출제됨', '직접 답'],
    ['삭제 이상',           '출제됨', '직접 답'],
    ['갱신 이상',           '출제됨', '직접 답'],
    ['Full Dependency',     '출제됨', '직접 답'],
    ['Partial Dependency',  '출제됨', '직접 답'],
    ['Transitive Dependency','출제됨','직접 답'],
])

add_table_section('3-4. 설계 / 모델링 / ER', ['항목', '상태', '비고'], [
    ['요구조건 분석',           '출제됨', '직접 답'],
    ['개념적 설계',             '출제됨', '직접 답'],
    ['논리적 설계',             '출제됨', '직접 답'],
    ['물리적 설계',             '출제됨', '직접 답'],
    ['구현',                   '출제됨', '직접 답'],
    ['연산',                   '출제됨', '직접 답'],
    ['구조',                   '출제됨', '직접 답'],
    ['제약조건',               '출제됨', '직접 답'],
    ['외부 스키마',             '출제됨', '직접 답'],
    ['개념 스키마',             '출제됨', '직접 답'],
    ['내부 스키마',             '출제됨', '직접 답'],
    ['개체집합',               '출제됨', '직접 답'],
    ['관계집합 속성',           '출제됨', '직접 답'],
    ['링크',                   '출제됨', '직접 답'],
    ['속성',                   '출제됨', '직접 답'],
    ['관계집합·개체집합 연결선','출제됨', '직접 답'],
])

add_table_section('3-5. 트랜잭션 / 회복 / 병행제어', ['항목', '상태', '비고'], [
    ['Atomicity',       '출제됨', '직접 답'],
    ['Consistency',     '미출제', '문제 지문에만 등장'],
    ['Isolation',       '출제됨', '직접 답'],
    ['Durability',      '미출제', '문제 지문에만 등장'],
    ['Rollback',        '출제됨', '직접 답'],
    ['REDO',            '출제됨', '직접 답'],
    ['UNDO',            '출제됨', '직접 답'],
    ['Locking',         '출제됨', '직접 답'],
    ['즉각 갱신 회복 기법','출제됨','직접 답'],
    ['낙관적 검증',     '미출제', ''],
    ['타임스탬프 순서', '미출제', ''],
    ['다중버전 동시성 제어','미출제',''],
])

add_table_section('3-6. 관계대수 / 관계해석 / 조인', ['항목', '상태', '비고'], [
    ['Select(σ)',           '출제됨', '직접 답'],
    ['Project(π)',          '출제됨', '직접 답'],
    ['Join(⋈)',             '출제됨', '직접 답'],
    ['Division(÷)',         '출제됨', '직접 답'],
    ['Union(∪)',            '출제됨', '직접 답'],
    ['Difference(-)',       '출제됨', '직접 답'],
    ['Cartesian Product(×)','출제됨','직접 답'],
    ['관계해석',            '출제됨', '직접 답'],
    ['세타 조인',           '출제됨', '직접 답'],
    ['동등 조인',           '출제됨', '직접 답'],
    ['자연 조인',           '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 4. 네트워크
# ════════════════════════════════════════════════════════════
add_section('4. 네트워크')

add_table_section('4-1. 기본 개념 / 계층', ['항목', '상태', '비고'], [
    ['구문(Syntax)',     '출제됨', '직접 답'],
    ['의미(Semantic)',   '출제됨', '직접 답'],
    ['타이밍(Timing)',   '출제됨', '직접 답'],
    ['프로토콜',         '출제됨', '직접 답'],
    ['물리 계층',        '출제됨', '직접 답'],
    ['데이터링크 계층',  '출제됨', '직접 답'],
    ['네트워크 계층',    '출제됨', '직접 답'],
    ['표현 계층',        '출제됨', '직접 답'],
    ['전송 계층',        '미출제', ''],
    ['세션 계층',        '미출제', ''],
    ['응용 계층',        '미출제', ''],
])

add_table_section('4-2. 주소 / 라우팅 / 교환', ['항목', '상태', '비고'], [
    ['RIP',             '출제됨', '직접 답'],
    ['OSPF',            '출제됨', '직접 답'],
    ['BGP',             '출제됨', '직접 답'],
    ['IGP',             '출제됨', '직접 답'],
    ['EGP',             '출제됨', '직접 답'],
    ['IPv6',            '출제됨', '직접 답'],
    ['128비트',         '출제됨', '직접 답'],
    ['8비트',           '출제됨', '직접 답'],
    ['가상회선',         '출제됨', '직접 답'],
    ['데이터그램',       '출제됨', '직접 답'],
    ['ATM',             '출제됨', '직접 답'],
    ['NAT',             '출제됨', '직접 답'],
    ['VPN',             '출제됨', '직접 답'],
    ['네트워크 주소',    '출제됨', '직접 답'],
    ['브로드캐스트 주소','출제됨', '직접 답'],
    ['호스트 수 62',     '출제됨', '직접 답'],
])

add_table_section('4-3. 프로토콜 / 웹 / 통신', ['항목', '상태', '비고'], [
    ['ARP',             '출제됨', '직접 답'],
    ['RARP',            '출제됨', '직접 답'],
    ['ICMP',            '출제됨', '직접 답'],
    ['AJAX',            '출제됨', '직접 답'],
    ['JSON',            '출제됨', '직접 답'],
    ['SOAP',            '출제됨', '직접 답'],
    ['WSDL',            '출제됨', '직접 답'],
    ['XML',             '출제됨', '직접 답'],
    ['HTTP',            '출제됨', '직접 답'],
    ['HTML',            '출제됨', '직접 답'],
    ['Hypertext',       '출제됨', '직접 답'],
    ['SSH',             '출제됨', '직접 답'],
    ['L2TP',            '출제됨', '직접 답'],
    ['Ad-hoc Network',  '출제됨', '직접 답'],
    ['IaaS',            '출제됨', '직접 답'],
    ['PaaS',            '출제됨', '직접 답'],
    ['SaaS',            '출제됨', '직접 답'],
    ['정보 프레임',      '출제됨', '직접 답'],
    ['감독 프레임',      '출제됨', '직접 답'],
    ['비번호 프레임',    '출제됨', '직접 답'],
    ['동기균형모드',     '출제됨', '직접 답'],
    ['비동기응답모드',   '출제됨', '직접 답'],
    ['정상응답모드',     '미출제', ''],
    ['HTTPS',           '미출제', ''],
    ['TCP',             '미출제', ''],
    ['UDP',             '미출제', ''],
    ['FTP',             '미출제', ''],
    ['SMTP',            '미출제', ''],
    ['POP3',            '미출제', ''],
    ['IMAP',            '미출제', ''],
])

add_table_section('4-4. 기타 네트워크', ['항목', '상태', '비고'], [
    ['URL 구조 번호 매핑','출제됨', '직접 답'],
    ['Hamming',          '출제됨', '직접 답'],
    ['FEC',              '출제됨', '직접 답'],
    ['BEC',              '출제됨', '직접 답'],
    ['Parity',           '출제됨', '직접 답'],
    ['CRC',              '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 5. 보안
# ════════════════════════════════════════════════════════════
add_section('5. 보안')

add_table_section('5-1. 암호 / 인증 / 접근통제', ['항목', '상태', '비고'], [
    ['AES',             '출제됨', '직접 답'],
    ['DES',             '출제됨', '직접 답'],
    ['IDEA',            '출제됨', '직접 답'],
    ['SKIPJACK',        '출제됨', '직접 답'],
    ['TKIP',            '출제됨', '직접 답'],
    ['RSA',             '출제됨', '직접 답'],
    ['ECC',             '출제됨', '직접 답'],
    ['ARIA',            '출제됨', '직접 답'],
    ['SEED',            '출제됨', '직접 답'],
    ['MD5',             '출제됨', '직접 답'],
    ['해시(Hash)',       '출제됨', '직접 답'],
    ['IPSec',           '출제됨', '직접 답'],
    ['OAuth',           '출제됨', '직접 답'],
    ['OTP',             '출제됨', '직접 답'],
    ['SSO',             '출제됨', '직접 답'],
    ['MAC',             '출제됨', '직접 답'],
    ['RBAC',            '출제됨', '직접 답'],
    ['DAC',             '출제됨', '직접 답'],
    ['Authentication',  '출제됨', '직접 답'],
    ['Authorization',   '출제됨', '직접 답'],
    ['Accounting',      '출제됨', '직접 답'],
    ['SHA-1',           '미출제', ''],
    ['SHA-256',         '미출제', ''],
    ['3DES',            '미출제', ''],
    ['RC4',             '미출제', ''],
    ['ElGamal',         '미출제', ''],
    ['Diffie-Hellman',  '미출제', ''],
])

add_table_section('5-2. 공격 / 악성코드 / 보안용어', ['항목', '상태', '비고'], [
    ['LAND Attack',     '출제됨', '직접 답'],
    ['Sniffing',        '출제됨', '직접 답'],
    ['SQL Injection',   '출제됨', '직접 답'],
    ['Session Hijacking','출제됨','직접 답'],
    ['ARP 스푸핑',      '출제됨', '직접 답'],
    ['Watering Hole',   '출제됨', '직접 답'],
    ['Rootkit',         '출제됨', '직접 답'],
    ['스캐어웨어',       '출제됨', '직접 답'],
    ['웜',              '출제됨', '직접 답'],
    ['트로이목마',       '출제됨', '직접 답'],
    ['바이러스',         '출제됨', '직접 답'],
    ['사회공학',         '출제됨', '직접 답'],
    ['APT',             '출제됨', '직접 답'],
    ['스머프',           '출제됨', '직접 답'],
    ['SYN Flooding',    '출제됨', '직접 답'],
    ['템퍼프루핑',       '출제됨', '직접 답'],
    ['ISMS',            '출제됨', '직접 답'],
    ['SIEM',            '출제됨', '직접 답'],
    ['가용성',           '출제됨', '직접 답'],
    ['타이포스쿼팅',     '출제됨', '직접 답'],
    ['TrustZone',       '출제됨', '직접 답'],
    ['다크 데이터',      '출제됨', '직접 답'],
    ['XSS',             '미출제', ''],
    ['CSRF',            '미출제', ''],
    ['Replay Attack',   '미출제', ''],
    ['Evil Twin',       '미출제', ''],
    ['Ransomware',      '미출제', ''],
    ['기밀성',           '미출제', ''],
    ['무결성',           '미출제', ''],
])


# ════════════════════════════════════════════════════════════
# 6. 운영체제 / 시스템 / 신기술
# ════════════════════════════════════════════════════════════
add_section('6. 운영체제 / 시스템 / 신기술')

add_table_section('6-1. 운영체제', ['항목', '상태', '비고'], [
    ['HRN',             '출제됨', '직접 답'],
    ['SJF',             '출제됨', '직접 답'],
    ['SRT',             '출제됨', '직접 답'],
    ['RR',              '출제됨', '직접 답'],
    ['IPC',             '출제됨', '직접 답'],
    ['준비 상태',        '출제됨', '직접 답'],
    ['실행 상태',        '출제됨', '직접 답'],
    ['대기 상태',        '출제됨', '직접 답'],
    ['유닉스',           '출제됨', '직접 답'],
    ['안드로이드',       '출제됨', '직접 답'],
    ['LRU',             '출제됨', '직접 답'],
    ['LFU',             '출제됨', '직접 답'],
    ['FCFS',            '미출제', ''],
    ['우선순위 스케줄링','미출제', ''],
    ['다단계 큐',        '미출제', ''],
    ['다단계 피드백 큐', '미출제', ''],
    ['FIFO',            '미출제', ''],
    ['OPT',             '미출제', ''],
    ['NUR',             '미출제', ''],
])

add_table_section('6-2. 시스템 / 신기술', ['항목', '상태', '비고'], [
    ['RTO',     '출제됨', '직접 답'],
    ['RAID 0',  '출제됨', '직접 답'],
    ['블록체인', '출제됨', '직접 답'],
    ['하둡',    '출제됨', '직접 답'],
    ['LOD',     '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 7. 프로그래밍 언어 활용
# ════════════════════════════════════════════════════════════
add_section('7. 프로그래밍 언어 활용')

p = doc.add_paragraph()
run = p.add_run('※ 프로그래밍은 코드 출력형이 대부분이므로, 직접 답으로 나온 개념/키워드 중심으로만 정리.')
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
p.paragraph_format.space_after = Pt(4)

add_table_section('7-1. Java', ['항목', '상태', '비고'], [
    ['static',     '출제됨', '직접 답'],
    ['super',      '출제됨', '직접 답'],
    ['implements', '출제됨', '직접 답'],
    ['new',        '출제됨', '직접 답'],
])

add_table_section('7-2. C / C++', ['항목', '상태', '비고'], [
    ['생성자(Constructor)',          '출제됨', 'C++ 문항 직접 답'],
    ['구조체 포인터 연산자 (->)',    '출제됨', '직접 답'],
])

add_table_section('7-3. Python', ['항목', '상태', '비고'], [
    ['extend',  '출제됨', '직접 답'],
    ['pop',     '출제됨', '직접 답'],
    ['reverse', '출제됨', '직접 답'],
    ['split',   '출제됨', '직접 답'],
])


# ════════════════════════════════════════════════════════════
# 8. 공식 12개 범위 최종 체크
# ════════════════════════════════════════════════════════════
add_section('8. 공식 12개 범위 최종 체크')

add_table_section('', ['공식 범위', '상태', '판단'], [
    ['현행 시스템 분석 및 요구사항 확인', '약출제', '기능/비기능 요구사항 정도만 직접 확인'],
    ['데이터 입출력 구현',               '출제됨', 'DB 전반 다수'],
    ['통합 구현',                        '약출제', 'EAI, 연계 관련 일부만 직접 확인'],
    ['서버프로그램 구현',                '약출제', '정적/동적 분석, JUnit, 형상관리 정도'],
    ['인터페이스 구현',                  '약출제', 'SOAP/WSDL/AJAX, UI/NUI, implements 정도'],
    ['화면 설계',                        '약출제', 'UI/UX/GUI/NUI, 유효성/직관성 정도'],
    ['애플리케이션 테스트',              '출제됨', '테스트 기법/커버리지/V모델 다수'],
    ['SQL 응용',                         '출제됨', '강출제'],
    ['소프트웨어 개발 보안 구축',        '출제됨', '강출제'],
    ['프로그래밍 언어 활용',             '출제됨', '강출제'],
    ['응용 SW 기초기술 활용',            '출제됨', '네트워크/운영체제/시스템 다수'],
    ['제품 소프트웨어 패키징',           '약출제', '릴리즈 노트 헤더 정도'],
])


# ════════════════════════════════════════════════════════════
# 9. 미출제 핵심 요약
# ════════════════════════════════════════════════════════════
add_section('9. 미출제 핵심 요약')

subsections = {
    '9-1. 소프트웨어공학': [
        'Strategy', 'Decorator', 'Composite', 'Facade', 'Builder', 'Prototype',
        '자료 결합도', '외부 결합도', '논리적 응집도',
        '기본 경로 커버리지', '다중 조건 커버리지', '상태 전이 테스트',
        '결정 테이블 테스트', '비교 테스트', '분류 트리 테스트',
        '페어와이즈 테스트', '유스케이스 테스트',
        '합성 관계', '객체 다이어그램', '컴포넌트 다이어그램',
        '배치 다이어그램', '복합체 구조 다이어그램',
        '유스케이스 다이어그램', '활동 다이어그램', '시퀀스 다이어그램',
    ],
    '9-2. SQL / DB': [
        'REVOKE', 'ASC', 'INNER JOIN', 'NATURAL JOIN',
        'LEFT OUTER JOIN', 'RIGHT OUTER JOIN', 'INTERSECT', 'MINUS',
        '기본키', '1NF', 'BCNF', '4NF', '5NF',
        '낙관적 검증', '타임스탬프 순서', '다중버전 동시성 제어',
    ],
    '9-3. 네트워크 / 보안 / OS': [
        '정상응답모드(NRM)', 'HTTPS', 'TCP', 'UDP', 'FTP',
        'SMTP', 'POP3', 'IMAP',
        'SHA-1', 'SHA-256', '3DES', 'RC4', 'ElGamal', 'Diffie-Hellman',
        'XSS', 'CSRF', 'Replay Attack', 'Evil Twin', 'Ransomware',
        '기밀성', '무결성',
        'FCFS', '우선순위 스케줄링', '다단계 큐', '다단계 피드백 큐',
        'FIFO', 'OPT', 'NUR',
    ],
}

for sub_title, items in subsections.items():
    add_section(sub_title, level=2)
    # 3열 그리드로 표시
    cols = 3
    rows_data = []
    for i in range(0, len(items), cols):
        chunk = items[i:i+cols]
        while len(chunk) < cols:
            chunk.append('')
        rows_data.append(chunk)

    tbl = doc.add_table(rows=len(rows_data), cols=cols)
    tbl.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows_data):
        tr = tbl.rows[r_idx]
        is_even = (r_idx % 2 == 1)
        for c_idx, text in enumerate(row_data):
            cell = tr.cells[c_idx]
            cell.width = Cm(5.3)
            set_cell_border(cell, 'DDDDDD', 4)
            set_cell_margin(cell, 60, 60, 120, 120)
            if text:
                set_cell_bg(cell, C_RED_BG if not is_even else RGBColor(0xFF, 0xD7, 0xD9))
            else:
                set_cell_bg(cell, RGBColor(0xF9, 0xF9, 0xF9))
            p   = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            if text:
                run.font.color.rgb = C_RED_TXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── 저장 ──────────────────────────────────────────────────
out_path = '/home/jammy/projects/coding-test-practice/정보처리기사_실기_출제현황_최종판.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
