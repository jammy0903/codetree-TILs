from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


C_GREEN_BG = RGBColor(0xC6, 0xEF, 0xCE)
C_GREEN_TXT = RGBColor(0x27, 0x6F, 0x38)
C_YELLOW_BG = RGBColor(0xFF, 0xEB, 0x9C)
C_YELLOW_TXT = RGBColor(0x9C, 0x65, 0x00)
C_RED_BG = RGBColor(0xFF, 0xC7, 0xCE)
C_RED_TXT = RGBColor(0x9C, 0x00, 0x06)
C_HEADER_BG = RGBColor(0x26, 0x3D, 0x5E)
C_HEADER_TXT = RGBColor(0xFF, 0xFF, 0xFF)
C_ROW_ALT = RGBColor(0xF2, 0xF2, 0xF2)
C_TITLE_TXT = RGBColor(0x1F, 0x39, 0x64)


def set_cell_bg(cell, rgb: RGBColor):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tc_pr.append(shd)


def set_cell_border(cell, color="BBBBBB", sz=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        borders.append(border)
    tc_pr.append(borders)


def set_cell_margin(cell, top=60, bottom=60, left=100, right=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = OxmlElement("w:tcMar")
    for side, value in (
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ):
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")
        margin.append(elem)
    tc_pr.append(margin)


def status_colors(status: str):
    if status == "출제됨":
        return C_GREEN_BG, C_GREEN_TXT
    if status == "약출제":
        return C_YELLOW_BG, C_YELLOW_TXT
    if status == "미출제":
        return C_RED_BG, C_RED_TXT
    return None, None


def normalize_md_cell(text: str) -> str:
    text = text.strip()
    if text.startswith("`") and text.endswith("`") and len(text) >= 2:
        text = text[1:-1]
    return text


def parse_markdown(md_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = ""
    elements = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            if not title:
                title = stripped[2:].strip()
            i += 1
            continue

        if stripped.startswith("## "):
            elements.append(("heading1", stripped[3:].strip()))
            i += 1
            continue

        if stripped.startswith("### "):
            elements.append(("heading2", stripped[4:].strip()))
            i += 1
            continue

        if stripped.startswith(">"):
            text = stripped[1:].strip()
            if text and not text.startswith("- ") and "판정 기준" not in text and "기준 파일" not in text and "편집 원칙" not in text:
                elements.append(("paragraph", text))
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|---"):
            header = [normalize_md_cell(cell) for cell in stripped.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines):
                row_line = lines[i].strip()
                if not row_line.startswith("|"):
                    break
                row = [normalize_md_cell(cell) for cell in row_line.strip("|").split("|")]
                rows.append(row)
                i += 1
            elements.append(("table", header, rows))
            continue

        if stripped.startswith("- "):
            elements.append(("bullet", stripped[2:].strip()))
            i += 1
            continue

        elements.append(("paragraph", stripped))
        i += 1

    return title, elements


def add_title(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = C_TITLE_TXT
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)


def add_legend(doc: Document):
    items = [
        ("출제됨", C_GREEN_BG, C_GREEN_TXT, "정답 칸에 직접 등장"),
        ("약출제", C_YELLOW_BG, C_YELLOW_TXT, "문제/보기/요약/해설에서만 확인"),
        ("미출제", C_RED_BG, C_RED_TXT, "보유 기출 파일 기준 직접 확인되지 않음"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Normal Table"
    for i, (label, bg, fg, desc) in enumerate(items):
        cell = table.rows[0].cells[i]
        cell.width = Cm(5.3)
        set_cell_bg(cell, bg)
        set_cell_border(cell, "AAAAAA", 4)
        set_cell_margin(cell, 80, 80, 120, 120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = fg
        r2 = p.add_run(f"\n{desc}")
        r2.font.size = Pt(8)
        r2.font.color.rgb = fg
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = C_TITLE_TXT
    if level == 1:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)


def add_paragraph(doc: Document, text: str, bullet: bool = False):
    p = doc.add_paragraph(style="Normal")
    if bullet:
        p.style = "List Bullet"
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


def column_widths(col_count: int):
    if col_count == 2:
        return [Cm(7.0), Cm(9.0)]
    if col_count == 3:
        return [Cm(7.0), Cm(2.5), Cm(6.5)]
    if col_count == 4:
        return [Cm(4.2), Cm(2.1), Cm(8.3), Cm(3.4)]
    return [Cm(4.5)] * col_count


def add_table(doc: Document, headers, rows):
    col_n = len(headers)
    widths = column_widths(col_n)
    table = doc.add_table(rows=1 + len(rows), cols=col_n)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.width = widths[i]
        set_cell_bg(cell, C_HEADER_BG)
        set_cell_border(cell, "FFFFFF", 6)
        set_cell_margin(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = C_HEADER_TXT

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        is_even = r_idx % 2 == 1
        for c_idx in range(col_n):
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = row.cells[c_idx]
            cell.width = widths[c_idx]
            set_cell_border(cell, "DDDDDD", 4)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(9)
            if headers[c_idx] == "상태":
                bg, fg = status_colors(text)
                if bg:
                    set_cell_bg(cell, bg)
                    run.font.color.rgb = fg
                    run.font.bold = True
                elif is_even:
                    set_cell_bg(cell, C_ROW_ALT)
            elif is_even:
                set_cell_bg(cell, C_ROW_ALT)


def build_docx(md_path: Path, out_path: Path):
    title, elements = parse_markdown(md_path)
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_title(doc, title or md_path.stem)
    add_legend(doc)

    for element in elements:
        kind = element[0]
        if kind == "heading1":
            add_heading(doc, element[1], level=1)
        elif kind == "heading2":
            add_heading(doc, element[1], level=2)
        elif kind == "paragraph":
            add_paragraph(doc, element[1])
        elif kind == "bullet":
            add_paragraph(doc, element[1], bullet=True)
        elif kind == "table":
            _, headers, rows = element
            add_table(doc, headers, rows)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.save(out_path)


if __name__ == "__main__":
    base = Path("/home/jammy/projects/coding-test-practice")
    build_docx(
        base / "정보처리기사_실기_출제현황_뜻포함.md",
        base / "정보처리기사_실기_출제현황_뜻포함.docx",
    )
