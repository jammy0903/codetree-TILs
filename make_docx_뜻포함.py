"""
정보처리기사_실기_출제현황_뜻포함.md → docx 변환기.
가독성 최우선 + 가로(Landscape) 레이아웃 + 행별 상태 색상.
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path('/home/jammy/projects/coding-test-practice/정보처리기사_실기_출제현황_뜻포함.md')
OUT = Path('/home/jammy/projects/coding-test-practice/정보처리기사_실기_출제현황_뜻포함.docx')


# ─── 색상 팔레트 ───────────────────────────────────────────
C_GREEN_BG   = RGBColor(0xC6, 0xEF, 0xCE)
C_GREEN_TXT  = RGBColor(0x1E, 0x60, 0x30)
C_YELLOW_BG  = RGBColor(0xFF, 0xEB, 0x9C)
C_YELLOW_TXT = RGBColor(0x7A, 0x4D, 0x00)
C_RED_BG     = RGBColor(0xFF, 0xC7, 0xCE)
C_RED_TXT    = RGBColor(0x80, 0x00, 0x00)
C_HEADER_BG  = RGBColor(0x26, 0x3D, 0x5E)
C_HEADER_TXT = RGBColor(0xFF, 0xFF, 0xFF)
C_ROW_ALT    = RGBColor(0xF4, 0xF6, 0xFA)
C_TITLE_TXT  = RGBColor(0x1F, 0x39, 0x64)
C_ACCENT_BG  = RGBColor(0xFF, 0xF4, 0xCE)
C_ACCENT_TXT = RGBColor(0x7A, 0x4F, 0x01)


# ─── XML 헬퍼 ──────────────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_border(cell, color='DDDDDD', sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tcPr.append(borders)


def set_cell_margin(cell, top=70, bottom=70, left=110, right=110):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)


def row_status_colors(val: str):
    """
    행 전체에 적용할 (배경색, 텍스트색).
    '미출제' → 빨강, '약출제*' → 노랑, 그 외 값 있음 → 초록(출제됨).
    """
    v = val.strip()
    if not v:
        return None, None
    if v == '미출제':
        return C_RED_BG, C_RED_TXT
    if v.startswith('약출제'):
        return C_YELLOW_BG, C_YELLOW_TXT
    # 실제 회차 기입 or '출제됨'
    return C_GREEN_BG, C_GREEN_TXT


# ─── 문서 준비 (가로 방향) ────────────────────────────────
doc = Document()
section = doc.sections[0]

# A4 Landscape 설정
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width  = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = section.right_margin = Cm(1.8)
section.top_margin  = section.bottom_margin = Cm(1.5)
# 가용 폭 ≈ 29.7 - 3.6 = 26.1cm

# 기본 스타일 폰트 (한글 가독성)
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
style.font.size = Pt(10)


# ─── 타이틀 ───────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('정보처리기사 실기 출제현황 · 뜻 포함판')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = C_TITLE_TXT
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2020~2025 기출 분석 · 2026 쪽집게 전략 포함')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x66)
p.paragraph_format.space_after = Pt(8)


# ─── 범례 ─────────────────────────────────────────────────
legend_items = [
    ('출제됨 (회차 기입)', C_GREEN_BG, C_GREEN_TXT, '정답 칸에 직접 등장한 항목 — 행 전체 초록'),
    ('약출제',             C_YELLOW_BG, C_YELLOW_TXT, '문제/보기/해설에서만 확인 — 행 전체 노랑'),
    ('미출제',             C_RED_BG,    C_RED_TXT,   '보유 기출 기준 미확인 — 행 전체 빨강'),
]
leg = doc.add_table(rows=1, cols=3)
leg.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, bg, fg, desc) in enumerate(legend_items):
    cell = leg.rows[0].cells[i]
    cell.width = Cm(8.7)
    set_cell_bg(cell, bg)
    set_cell_border(cell, 'AAAAAA', 4)
    set_cell_margin(cell, 90, 90, 130, 130)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cp.add_run(label)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = fg
    r2 = cp.add_run(f'\n{desc}')
    r2.font.size = Pt(8)
    r2.font.color.rgb = fg

doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── 섹션 제목 ────────────────────────────────────────────
def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = C_TITLE_TXT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '263D5E')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x7C)


def add_para(text, size=10, bold=False, color=None, italic=False,
             space_before=0, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_callout(lines, bg=C_ACCENT_BG, fg=C_ACCENT_TXT, title=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Cm(26.1)
    set_cell_bg(cell, bg)
    set_cell_border(cell, 'E3C97A', 6)
    set_cell_margin(cell, 130, 130, 180, 180)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    if title:
        r = cp.add_run(title)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = fg
        cp.add_run('\n')
    for i, line in enumerate(lines):
        if i > 0 or title:
            cp.add_run('\n')
        r = cp.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = fg
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─── 표 빌더 ───────────────────────────────────────────────
def make_table(headers, rows, status_col=None):
    """
    status_col: '출제 이력' 또는 '상태' 컬럼 index.
    행 전체에 미출제/약출제/출제됨 색상 적용.
    총 가용 폭 26.1cm (A4 가로, 마진 1.8cm).
    """
    n = len(headers)

    # 4컬럼: 항목(4.5) / 출제이력(6.0) / 뜻설명(13.0) / 비고(2.6) = 26.1
    # 3컬럼: 영역(4.5) / 항목(6.0) / 뜻요약(15.6) = 26.1
    # 2컬럼: 영역(6.0) / 내용(20.1) = 26.1
    width_presets = {
        2: [Cm(6.0),  Cm(20.1)],
        3: [Cm(4.5),  Cm(6.0),  Cm(15.6)],
        4: [Cm(4.5),  Cm(6.0),  Cm(13.0), Cm(2.6)],
    }
    widths = width_presets.get(n, [Cm(26.1 / n)] * n)

    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    tblPr = tbl._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)

    # ── 헤더 행 ──
    hdr = tbl.rows[0]
    hdr.height = Cm(0.8)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = widths[i]
        set_cell_bg(cell, C_HEADER_BG)
        set_cell_border(cell, 'FFFFFF', 6)
        set_cell_margin(cell, 90, 90, 110, 110)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_HEADER_TXT

    # ── 데이터 행 ──
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        is_even = (ri % 2 == 1)

        # 행 상태 판단
        row_bg, status_txt = None, None
        if status_col is not None:
            val = row[status_col] if status_col < len(row) else ''
            row_bg, status_txt = row_status_colors(val)

        for ci, text in enumerate(row):
            cell = tr.cells[ci]
            cell.width = widths[ci]
            set_cell_border(cell, 'D8DCE6', 4)
            set_cell_margin(cell, 65, 65, 110, 110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            r = cp.add_run(text)
            r.font.size = Pt(9.5)

            if row_bg:
                # 행 전체에 상태 색상 적용
                set_cell_bg(cell, row_bg)
                if ci == status_col:
                    # 출제이력 셀: 굵게 + 진한 텍스트색
                    r.font.bold = True
                    r.font.color.rgb = status_txt
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif ci == 0:
                    r.font.bold = True
            else:
                # 상태 없는 표: 제브라 + 첫 열 굵게
                if ci == 0:
                    r.font.bold = True
                if is_even:
                    set_cell_bg(cell, C_ROW_ALT)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


# ─── 마크다운 파서 ─────────────────────────────────────────
def parse_md(text):
    lines = text.splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped.startswith('# '):
            blocks.append(('h1', stripped[2:].strip()))
            i += 1
            continue
        if stripped.startswith('## '):
            blocks.append(('h2', stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith('### '):
            blocks.append(('h3', stripped[4:].strip()))
            i += 1
            continue
        if re.match(r'^-{3,}$', stripped):
            blocks.append(('hr',))
            i += 1
            continue
        if stripped.startswith('>'):
            qlines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qlines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            blocks.append(('quote', qlines))
            continue
        # 표 감지
        if stripped.startswith('|') and i + 1 < len(lines) and re.match(r'^\s*\|?\s*:?-+', lines[i + 1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if len(cells) < len(header):
                    cells += [''] * (len(header) - len(cells))
                elif len(cells) > len(header):
                    cells = cells[:len(header)]
                rows.append(cells)
                i += 1
            blocks.append(('table', header, rows))
            continue
        if re.match(r'^\d+\.\s', stripped):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s', '', lines[i].strip()))
                i += 1
            blocks.append(('ol', items))
            continue
        if stripped.startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(('ul', items))
            continue
        blocks.append(('para', stripped))
        i += 1

    return blocks


# ─── 인라인 처리 ──────────────────────────────────────────
def clean_inline(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text


def render_inline_para(text, size=10, space_before=0, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'
        else:
            r = p.add_run(part)
        r.font.size = Pt(size)
    return p


def render_list(items, ordered=False):
    for idx, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        marker = f'{idx + 1}. ' if ordered else '• '
        r = p.add_run(marker)
        r.font.size = Pt(10)
        r.font.bold = ordered
        r.font.color.rgb = C_TITLE_TXT if ordered else RGBColor(0x55, 0x55, 0x66)
        parts = re.split(r'(\*\*.+?\*\*|`[^`]+`)', item)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2])
                r.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                r.font.name = 'Consolas'
            else:
                r = p.add_run(part)
            r.font.size = Pt(10)


# ─── 상태 컬럼 자동 감지 ─────────────────────────────────
def detect_status_col(headers):
    """'상태' 또는 '출제 이력' 컬럼 위치 반환."""
    for i, h in enumerate(headers):
        if h.strip() in ('상태', '출제 이력'):
            return i
    return None


# ─── 렌더링 ────────────────────────────────────────────────
md_text = SRC.read_text(encoding='utf-8')
blocks = parse_md(md_text)

skip_first_h1 = True

for b in blocks:
    kind = b[0]

    if kind == 'h1':
        if skip_first_h1:
            skip_first_h1 = False
            continue
        add_h1(clean_inline(b[1]))
    elif kind == 'h2':
        add_h1(clean_inline(b[1]))
    elif kind == 'h3':
        add_h2(clean_inline(b[1]))
    elif kind == 'hr':
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    elif kind == 'quote':
        add_callout(
            [clean_inline(l) for l in b[1] if l],
            bg=RGBColor(0xF5, 0xF7, 0xFB),
            fg=RGBColor(0x44, 0x4C, 0x5E),
        )
    elif kind == 'para':
        render_inline_para(b[1])
    elif kind == 'ol':
        render_list(b[1], ordered=True)
    elif kind == 'ul':
        render_list(b[1], ordered=False)
    elif kind == 'table':
        headers = [clean_inline(h) for h in b[1]]
        rows = [[clean_inline(c) for c in row] for row in b[2]]
        status_col = detect_status_col(headers)
        make_table(headers, rows, status_col=status_col)


# ─── 저장 ─────────────────────────────────────────────────
doc.save(OUT)
print(f'✔ 저장 완료: {OUT}')
